#!/usr/bin/env python3
"""
Enhanced Document Processor for ECHO
This module handles various document types including PDFs, Word documents, and other formats.
"""

import os
import json
import io
from typing import Optional, Dict, Any
import boto3

# Document processing libraries
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    print("⚠️  PyPDF2 not available. PDF processing will be limited.")

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("⚠️  python-docx not available. Word document processing will be limited.")

try:
    import pandas as pd
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False
    print("⚠️  pandas not available. Excel processing will be limited.")


class EnhancedDocumentProcessor:
    """Enhanced document processor for various file types"""

    def __init__(self, s3_client=None, bucket_name=None):
        self.s3_client = s3_client
        self.bucket_name = bucket_name
        self.supported_formats = {
            # Text formats
            '.txt': 'text',
            '.md': 'markdown',
            '.py': 'python',
            '.js': 'javascript',
            '.html': 'html',
            '.css': 'css',
            '.json': 'json',
            '.xml': 'xml',
            '.csv': 'csv',

            # Document formats
            '.pdf': 'pdf',
            '.docx': 'docx',
            '.doc': 'doc',
            '.xlsx': 'excel',
            '.xls': 'excel',

            # Code formats
            '.java': 'java',
            '.cpp': 'cpp',
            '.c': 'c',
            '.php': 'php',
            '.rb': 'ruby',
            '.go': 'go',
            '.rs': 'rust',
            '.swift': 'swift',
            '.kt': 'kotlin',
            '.scala': 'scala',

            # Data formats
            '.yaml': 'yaml',
            '.yml': 'yaml',
            '.toml': 'toml',
            '.ini': 'ini',
            '.conf': 'config',
        }

    def extract_content(self, key: str, content_type: str = '') -> Optional[Dict[str, Any]]:
        """
        Extract content from various file types

        Returns:
            Dict with keys: 'content', 'format', 'metadata', 'error'
        """
        try:
            # Get file extension
            file_extension = os.path.splitext(key)[1].lower()
            format_type = self.supported_formats.get(file_extension, 'unknown')

            # Get file from S3
            if self.s3_client and self.bucket_name:
                response = self.s3_client.get_object(
                    Bucket=self.bucket_name, Key=key)
                file_content = response['Body'].read()
                metadata = {
                    'size': response.get('ContentLength', 0),
                    'content_type': response.get('ContentType', content_type),
                    'last_modified': response.get('LastModified'),
                    'etag': response.get('ETag')
                }
            else:
                # For local testing
                with open(key, 'rb') as f:
                    file_content = f.read()
                metadata = {'size': len(file_content)}

            # Process based on format
            if format_type == 'pdf':
                return self._process_pdf(file_content, metadata)
            elif format_type == 'docx':
                return self._process_docx(file_content, metadata)
            elif format_type == 'excel':
                return self._process_excel(file_content, metadata)
            elif format_type in ['text', 'markdown', 'python', 'javascript', 'html', 'css', 'json', 'xml', 'csv']:
                return self._process_text(file_content, format_type, metadata)
            else:
                return self._process_unknown(file_content, format_type, metadata)

        except Exception as e:
            return {
                'content': None,
                'format': 'error',
                'metadata': {},
                'error': f"Error processing {key}: {str(e)}"
            }

    def _process_pdf(self, file_content: bytes, metadata: Dict) -> Dict[str, Any]:
        """Process PDF files"""
        if not PDF_AVAILABLE:
            return {
                'content': f"[PDF file - {metadata.get('size', 0)} bytes] PDF processing not available. Install PyPDF2: pip install PyPDF2",
                'format': 'pdf',
                'metadata': metadata,
                'error': 'PDF processing library not available'
            }

        try:
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)

            text_content = []
            total_pages = len(pdf_reader.pages)

            for page_num, page in enumerate(pdf_reader.pages, 1):
                page_text = page.extract_text()
                if page_text.strip():
                    text_content.append(
                        f"--- Page {page_num} ---\n{page_text}")

            content = "\n\n".join(text_content)

            # Add PDF metadata
            pdf_metadata = {
                'total_pages': total_pages,
                'pdf_info': pdf_reader.metadata if hasattr(pdf_reader, 'metadata') else {}
            }
            metadata.update(pdf_metadata)

            return {
                'content': content,
                'format': 'pdf',
                'metadata': metadata,
                'error': None
            }

        except Exception as e:
            return {
                'content': f"[PDF file - {metadata.get('size', 0)} bytes] Error extracting text: {str(e)}",
                'format': 'pdf',
                'metadata': metadata,
                'error': str(e)
            }

    def _process_docx(self, file_content: bytes, metadata: Dict) -> Dict[str, Any]:
        """Process Word documents"""
        if not DOCX_AVAILABLE:
            return {
                'content': f"[Word document - {metadata.get('size', 0)} bytes] Word processing not available. Install python-docx: pip install python-docx",
                'format': 'docx',
                'metadata': metadata,
                'error': 'Word processing library not available'
            }

        try:
            doc_file = io.BytesIO(file_content)
            doc = Document(doc_file)

            text_content = []

            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_content.append(para.text)

            # Extract tables
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells]
                    table_text.append(" | ".join(row_text))
                if table_text:
                    text_content.append("\n".join(table_text))

            content = "\n\n".join(text_content)

            # Add document metadata
            doc_metadata = {
                'paragraphs': len(doc.paragraphs),
                'tables': len(doc.tables),
                'sections': len(doc.sections)
            }
            metadata.update(doc_metadata)

            return {
                'content': content,
                'format': 'docx',
                'metadata': metadata,
                'error': None
            }

        except Exception as e:
            return {
                'content': f"[Word document - {metadata.get('size', 0)} bytes] Error extracting text: {str(e)}",
                'format': 'docx',
                'metadata': metadata,
                'error': str(e)
            }

    def _process_excel(self, file_content: bytes, metadata: Dict) -> Dict[str, Any]:
        """Process Excel files"""
        if not EXCEL_AVAILABLE:
            return {
                'content': f"[Excel file - {metadata.get('size', 0)} bytes] Excel processing not available. Install pandas: pip install pandas openpyxl",
                'format': 'excel',
                'metadata': metadata,
                'error': 'Excel processing library not available'
            }

        try:
            excel_file = io.BytesIO(file_content)

            # Read all sheets
            excel_data = pd.read_excel(excel_file, sheet_name=None)

            text_content = []

            for sheet_name, df in excel_data.items():
                if not df.empty:
                    text_content.append(f"--- Sheet: {sheet_name} ---")
                    text_content.append(
                        f"Dimensions: {df.shape[0]} rows x {df.shape[1]} columns")
                    text_content.append("Data:")
                    text_content.append(df.to_string(index=False, max_rows=50))
                    text_content.append("")

            content = "\n".join(text_content)

            # Add Excel metadata
            excel_metadata = {
                'sheets': len(excel_data),
                'sheet_names': list(excel_data.keys())
            }
            metadata.update(excel_metadata)

            return {
                'content': content,
                'format': 'excel',
                'metadata': metadata,
                'error': None
            }

        except Exception as e:
            return {
                'content': f"[Excel file - {metadata.get('size', 0)} bytes] Error extracting data: {str(e)}",
                'format': 'excel',
                'metadata': metadata,
                'error': str(e)
            }

    def _process_text(self, file_content: bytes, format_type: str, metadata: Dict) -> Dict[str, Any]:
        """Process text-based files"""
        try:
            content = file_content.decode('utf-8')

            # Handle JSON files
            if format_type == 'json':
                try:
                    json_data = json.loads(content)
                    content = json.dumps(json_data, indent=2)
                except:
                    pass  # Keep as plain text if JSON parsing fails

            return {
                'content': content,
                'format': format_type,
                'metadata': metadata,
                'error': None
            }

        except UnicodeDecodeError:
            # Try other encodings
            for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                try:
                    content = file_content.decode(encoding)
                    return {
                        'content': content,
                        'format': format_type,
                        'metadata': metadata,
                        'error': None
                    }
                except UnicodeDecodeError:
                    continue

            return {
                'content': f"[{format_type.upper()} file - {metadata.get('size', 0)} bytes] Unable to decode content",
                'format': format_type,
                'metadata': metadata,
                'error': 'Unicode decode error'
            }

    def _process_unknown(self, file_content: bytes, format_type: str, metadata: Dict) -> Dict[str, Any]:
        """Process unknown file types"""
        try:
            # Try to decode as text first
            content = file_content.decode('utf-8')
            return {
                'content': content[:2000] + "..." if len(content) > 2000 else content,
                'format': format_type,
                'metadata': metadata,
                'error': None
            }
        except UnicodeDecodeError:
            return {
                'content': f"[{format_type.upper()} file - {metadata.get('size', 0)} bytes] Binary file, cannot extract text content",
                'format': format_type,
                'metadata': metadata,
                'error': 'Binary file'
            }

    def get_supported_formats(self) -> Dict[str, str]:
        """Get list of supported file formats"""
        return self.supported_formats.copy()

    def is_format_supported(self, file_extension: str) -> bool:
        """Check if a file format is supported"""
        return file_extension.lower() in self.supported_formats

# Test function


def test_document_processor():
    """Test the document processor"""
    print("🧪 Testing Enhanced Document Processor")
    print("=" * 50)

    processor = EnhancedDocumentProcessor()

    print("📋 Supported Formats:")
    formats = processor.get_supported_formats()
    for ext, format_type in formats.items():
        print(f"   {ext} -> {format_type}")

    print(f"\n📊 Processing Libraries:")
    print(
        f"   PDF Processing: {'✅ Available' if PDF_AVAILABLE else '❌ Not Available'}")
    print(
        f"   Word Processing: {'✅ Available' if DOCX_AVAILABLE else '❌ Not Available'}")
    print(
        f"   Excel Processing: {'✅ Available' if EXCEL_AVAILABLE else '❌ Not Available'}")

    if not PDF_AVAILABLE or not DOCX_AVAILABLE or not EXCEL_AVAILABLE:
        print(f"\n💡 To enable full document processing, install:")
        if not PDF_AVAILABLE:
            print(f"   pip install PyPDF2")
        if not DOCX_AVAILABLE:
            print(f"   pip install python-docx")
        if not EXCEL_AVAILABLE:
            print(f"   pip install pandas openpyxl")


if __name__ == "__main__":
    test_document_processor()
